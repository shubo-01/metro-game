#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取花果山副本 PRD 和技术方案 Word 文档 - 直接以UTF-8写文件避免PowerShell管道编码问题"""

import io
import os
from docx import Document

DOCS = [
    ("PRD-花果山副本设计文档V3", r"E:\AIshubo\mePRD\word\寻仙 - 花果山副本设计文档 V3.docx"),
    ("技术方案-花果山副本",   r"E:\AIshubo\mePRD\word\寻仙 - 花果山副本技术方案.docx"),
]

OUT_DIR = r"e:\ziji-xiaochengxu\scripts"

def read_doc_to(buf, title, path):
    buf.write("=" * 100 + "\n")
    buf.write(f"# {title}\n")
    buf.write(f"文件: {path}\n")
    buf.write("=" * 100 + "\n")
    if not os.path.exists(path):
        buf.write(f"[错误] 文件不存在: {path}\n")
        return
    doc = Document(path)

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""
        buf.write(f"[P{i}][{style}] {text}\n")

    for ti, tbl in enumerate(doc.tables):
        buf.write(f"\n----- 表格{ti} 行数={len(tbl.rows)} -----\n")
        for ri, row in enumerate(tbl.rows):
            cells = [c.text.strip().replace("\n", "\\n") for c in row.cells]
            buf.write(f"  R{ri}: | " + " | ".join(cells) + " |\n")
    buf.write("\n")


# 分开两个文件写，方便阅读
for title, path in DOCS:
    fname = "hgs_prd.txt" if "PRD" in title else "hgs_tech.txt"
    outpath = os.path.join(OUT_DIR, fname)
    with io.open(outpath, "w", encoding="utf-8") as f:
        read_doc_to(f, title, path)
    print(f"[OK] 写入: {outpath} 大小={os.path.getsize(outpath)} bytes")
